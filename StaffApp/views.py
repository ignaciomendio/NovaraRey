from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from .models import *
from django.contrib.auth.models import User
from django.utils import timezone
from docx import Document
from django.http import HttpResponse
from AseguradoraApp.models import Cia, CiaTelContact, CiaMailContact
from MainApp.models import Address, AdressType
from django.urls import reverse
from CotizacionesApp.models import QuotRequest, Cotizacion
from EmisionesApp.models import Emision
from ClientesApp.models import ClientePersonaFisica, TarjetaCredito, Cliente, ClientePersonaJuridica

DASHBOARD_COLORS={
    "critic": "#e35d6a",
    "warning": "#ffc107",
    "OK": "#5ece4d"}

class Dashboard_Param():
    def __init__(self, name, value):
        self.name = name
        self.value = value

class Dashboard_data():
    title:str = ""
    params:list[Dashboard_Param] = []
    status:int = 0
    URL:str = ""

    def __init__(self, title, status,  URL):
        self.title = title
        self.status = status
        self.URL = URL

    def add_param(self, param:Dashboard_Param):
        self.params.append(param)

    def clean_params(self):
        self.params = []


@login_required(login_url='/login/login/')
def vista_dashboard(request):

    ds_context: list[Dashboard_data] = []

    #Código para el dasboard de cumpleaños
    clientes_cumple = list(filter(lambda x: x.aviso_cumple(), ClientePersonaFisica.objects.all()))
    if len(clientes_cumple) > 0:
        status_tareas = DASHBOARD_COLORS["warning"]
        ds_tareas: Dashboard_data = Dashboard_data(
            title="Clientes",
            status=status_tareas, 
            URL= reverse('clientes_main'))
        ds_tareas.clean_params()
        for cliente in clientes_cumple:
            ds_tareas.add_param(Dashboard_Param(f"{cliente.nombre} {cliente.apellido} cumple años el", f"{cliente.fecha_nacimiento.strftime('%d/%m')}"))
        ds_context.append(ds_tareas)


    #Código para el dasboard de tarjetas de crédito
    tarjetas_a_vencer = list(filter(lambda x: x.avisar_vencimiento(), TarjetaCredito.objects.all()))
    if len(tarjetas_a_vencer) > 0:
        status_tareas = DASHBOARD_COLORS["warning"]
        ds_tareas: Dashboard_data = Dashboard_data(
            title="Tarjetas",
            status=status_tareas, 
            URL= reverse('clientes_main'))
        ds_tareas.clean_params()
        for tarjeta in tarjetas_a_vencer:    
            if ClientePersonaFisica.objects.filter(id=tarjeta.Cliente.id).exists():
                cliente = ClientePersonaFisica.objects.filter(id=tarjeta.Cliente.id)
                nombre_cte = f'{cliente[0].nombre} {cliente[0].apellido}'
            else:
                cliente = ClientePersonaJuridica.objects.filter(id=tarjeta.Cliente.id)
                nombre_cte = f'{cliente[0].razon_social}'            
            ds_tareas.add_param(Dashboard_Param(f"{tarjeta.tarjeta} Nro. {tarjeta.numero_tarjeta} de {nombre_cte}", f"Vence: {tarjeta.vencimiento}"))
        ds_context.append(ds_tareas)


    #Código para el dasboard de Tareas
    tareas = Tarea.objects.filter(status="P")
    n_pend = len(tareas)
    n_venc = 0

    for tarea in tareas:
        if tarea.vencida():
            n_venc += 1

    status_tareas = DASHBOARD_COLORS["OK"]
    if n_pend > 0:
        status_tareas = DASHBOARD_COLORS["warning"]
    if n_venc > 0:
        status_tareas = DASHBOARD_COLORS["critic"]

    ds_tareas: Dashboard_data = Dashboard_data(
        title=Tarea._meta.verbose_name_plural,
        status=status_tareas, 
        URL= reverse('tareas_main'))
    ds_tareas.clean_params()
    ds_tareas.add_param(Dashboard_Param("Tareas Pendientes", n_pend))
    ds_tareas.add_param(Dashboard_Param("Tareas Vencidas", n_venc))

    ds_context.append(ds_tareas)

    #Código para el dasboard de solicitudes de cotización
    solicitudes = QuotRequest.objects.filter(status="P")
    n_pend = len(solicitudes)
    n_venc = len(list(filter(lambda x: x.cotizar_vencida(), solicitudes)))
    status_tareas = DASHBOARD_COLORS["OK"]
    if n_pend > 0:
        status_tareas = DASHBOARD_COLORS["warning"]
    if n_venc > 0:
        status_tareas = DASHBOARD_COLORS["critic"]
    ds_tareas: Dashboard_data = Dashboard_data(
        title=QuotRequest._meta.verbose_name_plural,
        status=status_tareas, 
        URL= reverse('ver_solicitudes_cotizacion')+"?status=P")
    ds_tareas.clean_params()
    ds_tareas.add_param(Dashboard_Param("Solicitudes Pendientes", n_pend))
    ds_tareas.add_param(Dashboard_Param(f"Solicitudes Vencidas (más de {QuotRequest.DIAS_MAX_SOLICITUD} días)", n_venc))

    ds_context.append(ds_tareas)

    #Código para el dasboard de cotizaciones en preparacion
    cot_prep = Cotizacion.objects.filter(status="P")
    n_pend = len(cot_prep)
    n_venc = len(list(filter(lambda x: x.preparar_vencida(), cot_prep)))
    status_tareas = DASHBOARD_COLORS["OK"]
    if n_pend > 0:
        status_tareas = DASHBOARD_COLORS["warning"]
    if n_venc > 0:
        status_tareas = DASHBOARD_COLORS["critic"]
    ds_tareas: Dashboard_data = Dashboard_data(
        title="Cotizaciones en Preparacion",
        status=status_tareas, 
        URL= reverse('ver_cotizaciones')+"?status=P")
    ds_tareas.clean_params()
    ds_tareas.add_param(Dashboard_Param("Cotizaciones en Preparacion", n_pend))
    ds_tareas.add_param(Dashboard_Param(f"Cotizaciones en Preparacion Vencidas (más de {Cotizacion.DIAS_MAX_PREPARACION} días)", n_venc))

    ds_context.append(ds_tareas)

    #Código para el dasboard de cotizaciones enviadas
    cot_prep = Cotizacion.objects.filter(status="E")
    n_pend = len(cot_prep)
    n_venc = len(list(filter(lambda x: x.enviar_vencida(), cot_prep)))
    status_tareas = DASHBOARD_COLORS["OK"]
    if n_pend > 0:
        status_tareas = DASHBOARD_COLORS["warning"]
    if n_venc > 0:
        status_tareas = DASHBOARD_COLORS["critic"]
    ds_tareas: Dashboard_data = Dashboard_data(
        title="Cotizaciones sin Emitir",
        status=status_tareas, 
        URL= reverse('ver_cotizaciones')+"?status=E")
    ds_tareas.clean_params()
    ds_tareas.add_param(Dashboard_Param("Cotizaciones enviadas a clientes sin emitir", n_pend))
    ds_tareas.add_param(Dashboard_Param(f"Cotizaciones sin emitir Vencidas (más de {Cotizacion.DIAS_MAX_RESPUESTA_CTE} días)", n_venc))

    ds_context.append(ds_tareas)

    #Código para el dasboard de emisiones sin pólizas
    emisiones_sin_poliza = Emision.objects.filter(tiene_poliza=False)
    n_pend = len(emisiones_sin_poliza)
    n_venc = len(list(filter(lambda x: x.cotizar_vencida(), emisiones_sin_poliza)))

    status_tareas = DASHBOARD_COLORS["OK"]
    if n_pend > 0:
        status_tareas = DASHBOARD_COLORS["warning"]
    if n_venc > 0:
        status_tareas = DASHBOARD_COLORS["critic"]
    ds_tareas: Dashboard_data = Dashboard_data(
        title="Emisiones sin póliza",
        status=status_tareas, 
        URL= reverse('ver_emisiones')+"?filtro_poliza=SP")
    ds_tareas.clean_params()
    ds_tareas.add_param(Dashboard_Param("Emisiones sin poliza asignada", n_pend))
    ds_tareas.add_param(Dashboard_Param(f"Emisiones Vencidas (más de {Emision.MAX_DIAS_SIN_POLIZA} días)", n_venc))

    ds_context.append(ds_tareas)

    return render(request, 'StaffApp/dashboard.html',{
        "ds_context":ds_context,})

@login_required(login_url='/login/login/')
def vista_enter(request):
    if request.user.is_authenticated:
        # If the user is authenticated, redirect to the dashboard
        return redirect('dashboard')
    else:
        return redirect('login')
    
@login_required(login_url='/login/login/')
def vista_tareas_main(request):
    tareas = Tarea.objects.filter(status='P').order_by('fecha_vencimiento')
    return render(request, 'StaffApp/tareasMain.html', {'tareas': tareas})


@login_required(login_url='/login/login/')
def vista_tareas_create(request):
    if request.method == 'POST':
        titulo = request.POST.get('titulo')
        fecha_vencimiento = request.POST.get('fecha_vencimiento')
        status = 'P'
        
        tarea = Tarea(
            titulo=titulo,
            usuario_creacion=request.user,
            fecha_vencimiento=fecha_vencimiento,
            status=status,
        )
        tarea.save()
        return redirect('tareas_main')
    
    return render(request, 'StaffApp/tareasCreate.html')


@login_required(login_url='/login/login/')
def vista_tareas_edit(request, id): 
    tarea = Tarea.objects.get(id=id)
    
    if request.method == 'POST':

        accion = request.POST.get('accion')
        if accion == 'addcomment':
            comentario = request.POST.get('comentario')
            if comentario:
                usuario = request.user.username
                timestamp = timezone.localtime(timezone.now()).strftime('%d/%m/%Y %H:%M')
                tarea.notas += f"\n[{usuario}:{timestamp}]: {comentario}"
                tarea.save()
                return redirect('tareas_main')
            else:
                # If no comment is provided, redirect to the edit view
                return redirect('tareas_edit', id=id)

        elif accion == 'finalizar':
            tarea.status = 'F'
            tarea.usuario_finalizacion = request.user
            tarea.fecha_finalizacion = timezone.now()
            tarea.save()
            return redirect('tareas_main')

        elif accion == 'cancelar':
            tarea.status = 'C'
            tarea.usuario_finalizacion = request.user
            tarea.fecha_finalizacion = timezone.now()
            tarea.save()
            return redirect('tareas_main')
            
    return render(request, 'StaffApp/tareasEdit.html', {'tarea': tarea})

@login_required(login_url='/login/login/')
def vista_tareas_export(request):
    from django.http import HttpResponse
    import pandas as pd

    # Obtén todos los campos de Tarea como diccionarios
    tareas = Tarea.objects.all().values('id', 'titulo', 'notas', 'status', 'fecha_creacion', 'fecha_vencimiento',
    'usuario_creacion__username', 'usuario_finalizacion__username', 'fecha_finalizacion')

    for tarea in tareas:
        tarea['fecha_creacion'] = tarea['fecha_creacion'].replace(tzinfo=None)
        tarea['fecha_vencimiento'] = tarea['fecha_vencimiento'].replace(tzinfo=None)
        if tarea['fecha_finalizacion']:
            tarea['fecha_finalizacion'] = tarea['fecha_finalizacion'].replace(tzinfo=None)

    df = pd.DataFrame(list(tareas))
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename="tareas.xlsx"'
    
    with pd.ExcelWriter(response, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Tareas')

    return response

@login_required(login_url='/login/login/')
def vista_prestadores_main(request):
    rubrofilter = request.GET.get('rubrofilter', None)
    activofilter = request.GET.get('activofilter', None)
    nombrefilter = request.GET.get('nombrefilter', None)
    prestadores = Prestador.objects.all().order_by('rubro__nombre', 'nombre')
    if rubrofilter:
        prestadores = prestadores.filter(rubro__id=rubrofilter)
    if activofilter:
        if activofilter == '1':
            prestadores = prestadores.filter(activo=True)
        elif activofilter == '0':
            prestadores = prestadores.filter(activo=False)
    if nombrefilter:
        prestadores = prestadores.filter(nombre__icontains=nombrefilter)
    rubros = RubroPrestador.objects.all().order_by('nombre')
    
    return render(request, 'StaffApp/prestadoresMain.html', {'prestadores': prestadores, 'rubros': rubros})
    

@login_required(login_url='/login/login/')
def vista_prestadores_create(request):
    if request.method == 'POST':
        nombre = request.POST.get('nombre')
        rubro_id = request.POST.get('rubro')
        contacto = request.POST.get('contacto')
        telefono = request.POST.get('telefono')
        email = request.POST.get('email')
        direccion = request.POST.get('direccion')
        sitio_web = request.POST.get('sitio_web')
        observaciones = request.POST.get('observaciones')
        activo = True

        rubro = RubroPrestador.objects.get(id=rubro_id)
        
        prestador = Prestador(
            nombre=nombre,
            rubro=rubro,
            contacto=contacto,
            telefono=telefono,
            email=email,    
            direccion=direccion,
            sitio_web=sitio_web,
            observaciones=observaciones,
            activo = True
        )
        prestador.save()
        return redirect('prestadores_main')
    
    rubros = RubroPrestador.objects.all().order_by('nombre')
    return render(request, 'StaffApp/prestadoresCreate.html', {'rubros': rubros})

@login_required(login_url='/login/login/')
def vista_prestadores_edit(request, id):
    prestador = Prestador.objects.get(id=id)
    
    if request.method == 'POST':
        nombre = request.POST.get('nombre')
        rubro_id = request.POST.get('rubro')
        contacto = request.POST.get('contacto')
        telefono = request.POST.get('telefono')
        email = request.POST.get('email')
        direccion = request.POST.get('direccion')
        sitio_web = request.POST.get('sitio_web')
        observaciones = request.POST.get('observaciones')
        activo = (request.POST.get('activo','0') == '1')

        rubro = RubroPrestador.objects.get(id=rubro_id)
        
        prestador.nombre = nombre
        prestador.rubro = rubro
        prestador.contacto = contacto
        prestador.telefono = telefono   
        prestador.email = email
        prestador.activo = activo
        prestador.direccion = direccion
        prestador.sitio_web = sitio_web
        prestador.observaciones = observaciones
        prestador.save()
        
        return redirect('prestadores_main')
    
    rubros = RubroPrestador.objects.all().order_by('nombre')
    return render(request, 'StaffApp/prestadoresEdit.html', {'prestador': prestador, 'rubros': rubros})

@login_required(login_url='/login/login/')
def vista_prestadores_doc(request):
    if request.method == 'POST':
        ids = request.POST.getlist('prestadores')
        prestadores = Prestador.objects.filter(id__in=ids)
        doc = Document()
        doc.add_heading('Prestadores Seleccionados', 0)
        for p in prestadores:
            doc.add_heading(p.nombre, level=1)
            doc.add_paragraph(f'Rubro: {p.rubro.nombre}')
            doc.add_paragraph(f'Contacto: {p.contacto or ""}')
            doc.add_paragraph(f'Teléfono: {p.telefono or ""}')
            doc.add_paragraph(f'Email: {p.email or ""}')
            doc.add_paragraph(f'Dirección: {p.direccion or ""}')
            doc.add_paragraph(f'Sitio Web: {p.sitio_web or ""}')
            doc.add_paragraph(f'Observaciones: {p.observaciones or ""}')
            doc.add_paragraph('---')
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=prestadores.docx'
        doc.save(response)
        return response
    return redirect('prestadores_main')

@login_required(login_url='/login/login/')
def vista_aseguradoras_main(request):
    nombrefilter = request.GET.get('nombrefilter', None)
    activofilter = request.GET.get('activofilter', None)
    aseguradoras = Cia.objects.all().order_by('nombre',)
    if activofilter:
        if activofilter == '1':
            aseguradoras = aseguradoras.filter(activa=True)
        elif activofilter == '0':
            aseguradoras = aseguradoras.filter(activa=False)
    if nombrefilter:
        aseguradoras = aseguradoras.filter(nombre__icontains=nombrefilter)
    
    return render(request, 'StaffApp/aseguradorasMain.html', {'aseguradoras': aseguradoras})

@login_required(login_url='/login/login/')
def vista_aseguradoras_create(request):
    pass

@login_required(login_url='/login/login/')
def vista_aseguradoras_edit(request, id):
    cia = Cia.objects.get(id=id)
    direccion = Address.objects.filter(cia=cia).first()
    telefonos = CiaTelContact.objects.filter(cia=cia)
    mails = CiaMailContact.objects.filter(cia=cia)
    addtypes = AdressType.objects.all().order_by('tipo')  

    if request.method == 'POST':
        cia.nombre = request.POST.get('nombre')
        cia.cuit = request.POST.get('cuit')
        cia.activa = request.POST.get('activa') == 'on'
        cia.logoUrl = request.POST.get('logoUrl')
        cia.url = request.POST.get('url')
        cia.desc = request.POST.get('desc')
        direccion_id = request.POST.get('direccion')

        if direccion_id:
            cia.direccion = get_object_or_404(Address, id=direccion_id)

        cia.save()
        return redirect('cia_list')  # Asegurate de tener definida esta URL

    return render(request, 'StaffApp/aseguradorasEdit.html', {
        'cia': cia,
        'direccion': direccion,
        'telefonos': telefonos,
        'mails': mails,
        'addtypes': addtypes,
    })

@login_required(login_url='/login/login/')
def vista_aseguradoras_edit_telAdd(request, id):
    cia = Cia.objects.get(id=id)
    if request.method == 'POST':
        telefono = request.POST.get('numero')
        contacto = request.POST.get('contactotel')
        desc = request.POST.get('descripcion')

        if telefono and contacto:
            tel_contact = CiaTelContact(
                telefono=telefono,
                contacto=contacto,
                desc=desc,
                cia=cia
            )
            tel_contact.save()
            return redirect('aseguradoras_edit', id=id)

    return render(request, 'StaffApp/aseguradorasEditTelAdd.html', {'cia': cia})

@login_required(login_url='/login/login/')
def vista_aseguradoras_edit_mailAdd(request, id):
    cia = Cia.objects.get(id=id)
    if request.method == 'POST':
        mail = request.POST.get('email')
        contacto = request.POST.get('contactomail')
        desc = request.POST.get('descripcionMail')

        if mail and contacto:
            mail_contact = CiaMailContact(
                mail=mail,
                contacto=contacto,
                desc=desc,
                cia=cia
            )
            mail_contact.save()
            return redirect('aseguradoras_edit', id=id)

    return render(request, 'StaffApp/aseguradorasEditMailAdd.html', {'cia': cia})

@login_required(login_url='/login/login/')
def vista_aseguradoras_edit_telEdit(request, id):
    tel_contact = CiaTelContact.objects.get(id=id)
    cia = tel_contact.cia

    if request.method == 'POST':
        tel_contact.telefono = request.POST.get('numero')
        tel_contact.contacto = request.POST.get('contacto')
        tel_contact.desc = request.POST.get('descripcion')
        tel_contact.save()
        return redirect('aseguradoras_edit', id=cia.id)

    return render(request, 'StaffApp/aseguradorasEditTelEdit.html', {'tel_contact': tel_contact, 'cia': cia})

@login_required(login_url='/login/login/')
def vista_aseguradoras_edit_mailEdit(request, id):
    mail_contact = CiaMailContact.objects.get(id=id)
    cia = mail_contact.cia

    if request.method == 'POST':
        mail_contact.mail = request.POST.get('mail')
        mail_contact.contacto = request.POST.get('contactomail')
        mail_contact.desc = request.POST.get('descripcionmail')
        mail_contact.save()
        return redirect('aseguradoras_edit', id=cia.id)

    return render(request, 'StaffApp/aseguradorasEditMailEdit.html', {'mail_contact': mail_contact, 'cia': cia})

@login_required(login_url='/login/login/')
def vista_aseguradoras_edit_telDelete(request, id):
    tel_contact = CiaTelContact.objects.get(id=id)
    cia = tel_contact.cia
    tel_contact.delete()
    return redirect('aseguradoras_edit', id=cia.id)

@login_required(login_url='/login/login/')
def vista_aseguradoras_edit_mailDelete(request, id):

    mail_contact = CiaMailContact.objects.get(id=id)
    cia = mail_contact.cia
    mail_contact.delete()
    return redirect('aseguradoras_edit', id=cia.id)

